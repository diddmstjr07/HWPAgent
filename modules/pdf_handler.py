#!/usr/bin/env python3
"""
PDF 파일 생성 모듈 - DOCX 변환 방식
"""
from pathlib import Path
from typing import Optional, List, Dict, Any
import os
import subprocess
import sys
import shutil
import re
import io

from docx import Document
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm, inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from xml.sax.saxutils import escape
import requests

WORD_NAMESPACES = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
}
EMUS_PER_INCH = 914400
MAX_IMAGE_WIDTH = 5.5 * inch


class PDFHandler:
    """한글을 지원하는 PDF 문서 생성 (DOCX 변환 방식)"""
    
    def __init__(self, output_dir: str = "output"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        self.fonts_dir = self.output_dir / 'fonts'
        self.fonts_dir.mkdir(exist_ok=True)
        self._fallback_font: Optional[str] = None
    
    def convert_docx_to_pdf(
        self,
        docx_path: str,
        output_filename: Optional[str] = None,
        style_config: Optional[Dict[str, Any]] = None
    ) -> str:
        """
        DOCX 파일을 PDF로 변환
        
        Args:
            docx_path: DOCX 파일 경로
            output_filename: 출력 PDF 파일명
        
        Returns:
            생성된 PDF 파일 경로
        """
        docx_path = Path(docx_path)
        if not docx_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {docx_path}")
        
        if output_filename is None:
            output_filename = docx_path.stem + ".pdf"
        
        output_path = self.output_dir / output_filename
        
        print(f"[PDF] Converting {docx_path.name} to PDF...")
        
        try:
            # macOS: LibreOffice 사용
            if sys.platform == 'darwin':
                return self._convert_with_libreoffice(docx_path, output_path)
            # Windows: win32com (Microsoft Word) 사용
            elif sys.platform == 'win32':
                return self._convert_with_word(docx_path, output_path)
            # Linux: LibreOffice 사용
            else:
                return self._convert_with_libreoffice(docx_path, output_path)
        except FileNotFoundError as exc:
            # LibreOffice 미설치 시 경량 PDF 생성기로 폴백
            print(f"[PDF] LibreOffice not available: {exc}")
            print("[PDF] Falling back to lightweight PDF builder (rich text + inline images).")
            return self._convert_with_reportlab(docx_path, output_path, style_config or {})
    
    def _convert_with_libreoffice(self, docx_path: Path, output_path: Path) -> str:
        """
        LibreOffice를 사용하여 DOCX를 PDF로 변환 (macOS/Linux)
        """
        try:
            # LibreOffice 명령어 경로 찾기
            libreoffice_paths = [
                Path('/Applications/LibreOffice.app/Contents/MacOS/soffice'),
                Path('/usr/bin/libreoffice'),
            ]
            
            soffice_cmd = None
            for candidate in libreoffice_paths:
                if candidate.exists():
                    soffice_cmd = str(candidate)
                    break
            if not soffice_cmd:
                soffice_cmd = shutil.which('libreoffice')
            
            if not soffice_cmd:
                raise FileNotFoundError(
                    "LibreOffice not found. Please install:\n"
                    "  macOS: brew install --cask libreoffice\n"
                    "  Linux: sudo apt-get install libreoffice"
                )
            
            # LibreOffice로 변환
            cmd = [
                soffice_cmd,
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', str(self.output_dir),
                str(docx_path)
            ]
            
            print(f"[PDF] Running LibreOffice conversion...")
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
            
            if result.returncode != 0:
                raise Exception(f"LibreOffice conversion failed: {result.stderr}")
            
            # 생성된 PDF 파일 경로
            generated_pdf = self.output_dir / f"{docx_path.stem}.pdf"
            
            # 원하는 파일명으로 변경
            if generated_pdf != output_path:
                if output_path.exists():
                    output_path.unlink()
                generated_pdf.rename(output_path)
            
            # 임시 DOCX 파일 삭제
            if '_temp' in docx_path.name:
                docx_path.unlink()
                print(f"[PDF] Cleaned up temp DOCX: {docx_path.name}")
            
            print(f"[PDF] ✅ PDF created: {output_path.name}")
            return str(output_path)
            
        except subprocess.TimeoutExpired:
            raise Exception("PDF conversion timeout (30s)")
        except FileNotFoundError:
            raise
        except Exception as e:
            print(f"[PDF ERROR] {str(e)}")
            raise
    
    def _convert_with_word(self, docx_path: Path, output_path: Path) -> str:
        """
        Microsoft Word를 사용하여 DOCX를 PDF로 변환 (Windows)
        """
        try:
            import win32com.client
            
            word = win32com.client.Dispatch('Word.Application')
            word.Visible = False
            
            doc = word.Documents.Open(str(docx_path.absolute()))
            doc.SaveAs(str(output_path.absolute()), FileFormat=17)  # 17 = PDF
            doc.Close()
            word.Quit()
            
            # 임시 DOCX 파일 삭제
            if '_temp' in docx_path.name:
                docx_path.unlink()
                print(f"[PDF] Cleaned up temp DOCX: {docx_path.name}")
            
            print(f"[PDF] ✅ PDF created: {output_path.name}")
            return str(output_path)
            
        except Exception as e:
            print(f"[PDF ERROR] {str(e)}")
            raise

    def _convert_with_reportlab(self, docx_path: Path, output_path: Path, style_config: Dict[str, Any]) -> str:
        """LibreOffice 없이 간단한 텍스트 기반 PDF 생성"""
        try:
            document = Document(str(docx_path))
            story = self._build_story_from_docx(document)
            font_name = self._resolve_pdf_font_name(
                style_config.get('font_name'),
                style_config.get('font_file_path')
            )
            heading_font = self._resolve_pdf_font_name(
                style_config.get('heading_font_name'),
                style_config.get('heading_font_file_path')
            )
            title_font = self._resolve_pdf_font_name(
                style_config.get('title_font_name'),
                style_config.get('title_font_file_path')
            )
            body_size = style_config.get('font_size', 11)
            heading_sizes = {
                1: style_config.get('heading_level1_size', body_size + 3),
                2: style_config.get('heading_level2_size', body_size + 2),
                3: style_config.get('heading_level3_size', body_size + 1)
            }
            title_size = style_config.get('title_size', heading_sizes[1] + 2)
            line_spacing = style_config.get('line_spacing', 1.3)

            styles = getSampleStyleSheet()
            body_style = ParagraphStyle(
                'Body',
                parent=styles['BodyText'],
                fontName=font_name,
                fontSize=body_size,
                leading=body_size * line_spacing,
                spaceAfter=6
            )
            heading_styles = {
                1: ParagraphStyle('Heading1', parent=styles['Heading1'], fontName=heading_font,
                                  fontSize=heading_sizes[1], leading=heading_sizes[1] * 1.2,
                                  spaceBefore=12, spaceAfter=6),
                2: ParagraphStyle('Heading2', parent=styles['Heading2'], fontName=heading_font,
                                  fontSize=heading_sizes[2], leading=heading_sizes[2] * 1.15,
                                  spaceBefore=10, spaceAfter=4),
                3: ParagraphStyle('Heading3', parent=styles['Heading3'], fontName=heading_font,
                                  fontSize=heading_sizes[3], leading=heading_sizes[3] * 1.1,
                                  spaceBefore=8, spaceAfter=4)
            }
            title_style = ParagraphStyle('Title', parent=styles['Title'], fontName=title_font,
                                         fontSize=title_size, leading=title_size * 1.2,
                                         alignment=1, spaceAfter=12)

            pdf_story = []
            number_counter = 1
            for block in story:
                if block['type'] == 'spacer':
                    pdf_story.append(Spacer(1, 0.2 * cm))
                    continue
                text = escape(block['text']).replace('\n', '<br/>')
                if block['type'] == 'heading':
                    level = min(block.get('level', 1), 3)
                    pdf_story.append(Paragraph(text, heading_styles[level]))
                    number_counter = 1
                elif block['type'] == 'title':
                    pdf_story.append(Paragraph(text, title_style))
                    number_counter = 1
                elif block['type'] == 'bullet':
                    pdf_story.append(Paragraph(f"• {text}", body_style))
                    number_counter = 1
                elif block['type'] == 'number':
                    pdf_story.append(Paragraph(f"{number_counter}. {text}", body_style))
                    number_counter += 1
                elif block['type'] == 'image':
                    image_flowable = self._build_pdf_image(block)
                    if image_flowable:
                        pdf_story.append(image_flowable)
                        pdf_story.append(Spacer(1, 0.2 * cm))
                    number_counter = 1
                else:
                    pdf_story.append(Paragraph(text, body_style))
                    number_counter = 1

            if not pdf_story:
                pdf_story.append(Paragraph(' ', body_style))

            doc = SimpleDocTemplate(
                str(output_path),
                pagesize=A4,
                leftMargin=2 * cm,
                rightMargin=2 * cm,
                topMargin=2 * cm,
                bottomMargin=2 * cm
            )
            doc.build(pdf_story)
            if '_temp' in docx_path.name and docx_path.exists():
                docx_path.unlink()
            print(f"[PDF] ✅ Lightweight PDF created: {output_path.name}")
            return str(output_path)
        except Exception as exc:
            print(f"[PDF FALLBACK ERROR] {exc}")
            raise Exception('PDF 변환 도중 오류가 발생했습니다. LibreOffice 설치 후 다시 시도해주세요.')

    def _build_story_from_docx(self, document: Document) -> List[Dict[str, Any]]:
        blocks: List[Dict[str, Any]] = []
        spacer_count = 0
        for para in document.paragraphs:
            image_blocks = self._extract_images_from_paragraph(para, document)
            text = para.text.strip()
            style_name = (para.style.name if para.style else '').lower()
            if not text:
                if image_blocks:
                    spacer_count = 0
                    blocks.extend(image_blocks)
                    continue
                spacer_count += 1
                if spacer_count < 3:
                    blocks.append({'type': 'spacer'})
                continue
            spacer_count = 0

            if 'title' in style_name and 'heading' not in style_name:
                blocks.append({'type': 'title', 'text': text})
                continue

            heading_match = re.search(r'heading\s*(\d)', style_name)
            if heading_match:
                level = int(heading_match.group(1))
                blocks.append({'type': 'heading', 'text': text, 'level': level})
                continue

            if 'list bullet' in style_name:
                blocks.append({'type': 'bullet', 'text': text})
                continue

            if 'list number' in style_name:
                blocks.append({'type': 'number', 'text': text})
                continue

            blocks.append({'type': 'paragraph', 'text': text})
            if image_blocks:
                blocks.extend(image_blocks)

        for table in document.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    blocks.append({'type': 'paragraph', 'text': row_text})
        return blocks

    def _extract_images_from_paragraph(self, paragraph, document: Document) -> List[Dict[str, Any]]:
        blocks: List[Dict[str, Any]] = []
        drawings = paragraph._element.xpath('.//w:drawing')
        for drawing in drawings:
            blip_elems = drawing.xpath('.//a:blip')
            if not blip_elems:
                continue
            embed_rid = blip_elems[0].get(qn('r:embed'))
            if not embed_rid:
                continue
            image_part = document.part.related_parts.get(embed_rid)
            if not image_part:
                continue
            blob = image_part.blob
            extent = drawing.xpath('.//wp:extent')
            width_pt = MAX_IMAGE_WIDTH
            height_pt = MAX_IMAGE_WIDTH * 0.6
            if extent:
                try:
                    cx = int(extent[0].get('cx', 0))
                    cy = int(extent[0].get('cy', 0))
                except (TypeError, ValueError):
                    cx = cy = 0
                if cx > 0:
                    width_pt = min(MAX_IMAGE_WIDTH, (cx / EMUS_PER_INCH) * inch)
                if cy > 0:
                    if cx > 0:
                        aspect = cy / cx
                        height_pt = max(1, width_pt * aspect)
                    else:
                        height_pt = max(1, (cy / EMUS_PER_INCH) * inch)
            blocks.append({
                'type': 'image',
                'data': blob,
                'width': width_pt,
                'height': height_pt
            })
        return blocks

    def _resolve_pdf_font_name(self, preferred_name: Optional[str], font_path: Optional[str]) -> str:
        safe_path = None
        if font_path:
            safe_path = self._resolve_font_path(font_path)
        registered = self._register_font_from_path(safe_path, preferred_name)
        if registered:
            return registered
        if preferred_name:
            return preferred_name
        return self._get_default_pdf_font()

    def _get_default_pdf_font(self) -> str:
        if self._fallback_font:
            return self._fallback_font
        font_candidates = [
            ('AppleGothic', '/System/Library/Fonts/AppleGothic.ttf'),
            ('NanumGothic', '/usr/share/fonts/truetype/nanum/NanumGothic.ttf'),
            ('NanumSquare', '/usr/share/fonts/truetype/nanum/NanumSquareRoundR.ttf'),
            ('NotoSansKR', '/usr/share/fonts/opentype/noto/NotoSansKR-Regular.otf'),
        ]
        for font_name, font_path in font_candidates:
            font_file = Path(font_path)
            if font_file.exists():
                registered = self._register_font_from_path(font_file, font_name)
                if registered:
                    self._fallback_font = registered
                    return registered
        downloaded = self._download_default_font()
        if downloaded:
            registered = self._register_font_from_path(downloaded, 'NotoSansKR')
            if registered:
                self._fallback_font = registered
                return registered
        self._fallback_font = 'Helvetica'
        return self._fallback_font

    def _register_font_from_path(self, font_path: Optional[Path], font_name: Optional[str]) -> Optional[str]:
        if not font_path or not font_path.exists():
            return None
        target_name = font_name or font_path.stem
        try:
            if font_path.suffix.lower() == '.ttc':
                pdfmetrics.registerFont(TTFont(target_name, str(font_path), subfontIndex=0))
            else:
                pdfmetrics.registerFont(TTFont(target_name, str(font_path)))
            return target_name
        except Exception as exc:
            print(f"[PDF FONT] Failed to register {font_path}: {exc}")
            return None

    def _resolve_font_path(self, font_path: str) -> Optional[Path]:
        try:
            path = Path(font_path).expanduser().resolve()
            fonts_root = self.fonts_dir.resolve()
            path.relative_to(fonts_root)
            return path
        except Exception:
            return None

    def _download_default_font(self) -> Optional[Path]:
        target = self.fonts_dir / 'NotoSansKR-Regular.otf'
        if target.exists():
            return target
        url = "https://github.com/googlefonts/noto-cjk/raw/main/Sans/OTF/Korean/NotoSansKR-Regular.otf"
        try:
            response = requests.get(url, timeout=30)
            response.raise_for_status()
            target.write_bytes(response.content)
            return target
        except Exception as exc:
            print(f"[PDF FONT] Failed to download default font: {exc}")
            return None

    def _build_pdf_image(self, block: Dict[str, Any]) -> Optional[RLImage]:
        data = block.get('data')
        if not data:
            return None
        width = block.get('width') or MAX_IMAGE_WIDTH
        height = block.get('height') or width * 0.6
        width = min(width, MAX_IMAGE_WIDTH)
        if height <= 0:
            height = width * 0.6
        try:
            img_stream = io.BytesIO(data)
            return RLImage(img_stream, width=width, height=height)
        except Exception as exc:
            print(f"[PDF] Failed to build image block: {exc}")
            return None
