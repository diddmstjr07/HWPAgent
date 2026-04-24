"""문서 양식 파일 파서"""
from pathlib import Path
from typing import Optional
from html import escape
import shutil
import fitz  # PyMuPDF
from docx import Document
from legacy.modules.hwp_handler import HWPHandler

SUPPORTED_TEMPLATE_EXTENSIONS = {'.docx', '.hwp', '.hwpx', '.txt', '.md', '.pdf'}


def extract_template_text(file_path: Path) -> str:
    """주어진 파일에서 양식 텍스트를 추출"""
    suffix = file_path.suffix.lower()
    if suffix not in SUPPORTED_TEMPLATE_EXTENSIONS:
        raise ValueError('지원하지 않는 파일 형식입니다. (.docx, .hwp, .hwpx, .pdf, .txt, .md)')

    try:
        if suffix == '.pdf':
            return _extract_text_from_pdf(file_path)
        
        if suffix in {'.hwp', '.hwpx'}:
            handler = HWPHandler()
            return handler.read_hwp(str(file_path))
        
        if suffix in {'.txt', '.md'}:
            return file_path.read_text(encoding='utf-8', errors='ignore').strip()

        return _extract_text_from_docx_like(file_path, treat_as_docx=(suffix == '.docx'))
    except Exception as exc:
        raise ValueError(f'파일을 읽는 중 오류가 발생했습니다: {str(exc)}') from exc


def extract_template_html(file_path: Path, template_id: Optional[str] = None) -> str:
    """주어진 파일에서 양식 HTML을 추출"""
    suffix = file_path.suffix.lower()
    if suffix not in SUPPORTED_TEMPLATE_EXTENSIONS:
        raise ValueError('지원하지 않는 파일 형식입니다. (.docx, .hwp, .hwpx, .pdf, .txt, .md)')

    try:
        if suffix == '.pdf':
            text = _extract_text_from_pdf(file_path)
            return _wrap_text_as_html(text)

        if suffix in {'.hwp', '.hwpx'}:
            handler = HWPHandler()
            return handler.convert_hwp_to_html(str(file_path), template_id=template_id)

        if suffix in {'.txt', '.md'}:
            text = file_path.read_text(encoding='utf-8', errors='ignore').strip()
            return _wrap_text_as_html(text)

        text = _extract_text_from_docx_like(file_path, treat_as_docx=(suffix == '.docx'))
        return _wrap_text_as_html(text)
    except Exception as exc:
        raise ValueError(f'파일을 변환하는 중 오류가 발생했습니다: {str(exc)}') from exc


def _extract_text_from_pdf(file_path: Path) -> str:
    """PDF 파일에서 텍스트 추출"""
    text_content = []
    try:
        with fitz.open(file_path) as doc:
            for page in doc:
                text_content.append(page.get_text())
        return "\n".join(text_content).strip()
    except Exception as e:
        raise ValueError(f"PDF 파싱 실패: {str(e)}")


def _extract_text_from_docx_like(file_path: Path, treat_as_docx: bool = True) -> str:
    temp_docx: Optional[Path] = None
    docx_path = file_path

    if not treat_as_docx:
        temp_docx = file_path.with_suffix('.docx')
        shutil.copyfile(file_path, temp_docx)
        docx_path = temp_docx

    try:
        document = Document(str(docx_path))
        lines = [para.text.rstrip() for para in document.paragraphs]

        table_lines = []
        for table in document.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells).strip()
                if row_text:
                    table_lines.append(row_text)

        combined = '\n'.join(line for line in lines if line is not None)
        if table_lines:
            combined = f"{combined}\n" + '\n'.join(table_lines)
        return combined.strip()
    finally:
        if temp_docx and temp_docx.exists():
            temp_docx.unlink()


def _wrap_text_as_html(text: str) -> str:
    escaped = escape(text or "")
    paragraphs = []
    buffer = []
    for line in escaped.splitlines():
        if line.strip():
            buffer.append(line)
        else:
            if buffer:
                paragraphs.append("<p>" + "<br>".join(buffer) + "</p>")
                buffer = []
    if buffer:
        paragraphs.append("<p>" + "<br>".join(buffer) + "</p>")
    body_html = "\n".join(paragraphs) if paragraphs else "<p></p>"
    return f'<div class="template-doc">{body_html}</div>'
