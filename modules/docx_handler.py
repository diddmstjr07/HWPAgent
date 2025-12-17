"""
DOCX (워드) 파일 생성 모듈 - 한글에서도 열리는 완전한 문서
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import shutil
from typing import Optional, Dict, Any, List
import os
from datetime import datetime
import re

LATEX_SYMBOL_MAP = {
    r'\times': '×',
    r'\cdot': '·',
    r'\pm': '±',
    r'\mp': '∓',
    r'\leq': '≤',
    r'\geq': '≥',
    r'\neq': '≠',
    r'\approx': '≈',
    r'\sim': '∼',
    r'\infty': '∞',
    r'\partial': '∂',
    r'\nabla': '∇',
    r'\sum': '∑',
    r'\prod': '∏',
    r'\int': '∫',
    r'\oint': '∮',
    r'\propto': '∝',
    r'\forall': '∀',
    r'\exists': '∃',
    r'\cup': '∪',
    r'\cap': '∩',
    r'\subseteq': '⊆',
    r'\supseteq': '⊇',
    r'\subset': '⊂',
    r'\supset': '⊃',
    r'\in': '∈',
    r'\notin': '∉',
    r'\perp': '⊥',
    r'\angle': '∠'
}

LATEX_GREEK_MAP = {
    r'\alpha': 'α', r'\beta': 'β', r'\gamma': 'γ', r'\delta': 'δ', r'\epsilon': 'ε',
    r'\zeta': 'ζ', r'\eta': 'η', r'\theta': 'θ', r'\iota': 'ι', r'\kappa': 'κ',
    r'\lambda': 'λ', r'\mu': 'μ', r'\nu': 'ν', r'\xi': 'ξ', r'\pi': 'π',
    r'\rho': 'ρ', r'\sigma': 'σ', r'\tau': 'τ', r'\upsilon': 'υ', r'\phi': 'φ',
    r'\chi': 'χ', r'\psi': 'ψ', r'\omega': 'ω',
    r'\Gamma': 'Γ', r'\Delta': 'Δ', r'\Theta': 'Θ', r'\Lambda': 'Λ', r'\Xi': 'Ξ',
    r'\Pi': 'Π', r'\Sigma': 'Σ', r'\Phi': 'Φ', r'\Psi': 'Ψ', r'\Omega': 'Ω'
}


class DOCXHandler:
    """완전한 서식을 갖춘 DOCX 문서 생성"""
    
    def __init__(self, output_dir: str = "output"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
    
    def create_document(
        self,
        title: str,
        content: str,
        style_config: Optional[Dict[str, Any]] = None,
        filename: Optional[str] = None,
        images: Optional[List[str]] = None,
        template_path: Optional[str] = None
    ) -> str:
        """
        완전한 서식을 갖춘 워드 문서 생성
        
        Args:
            title: 문서 제목
            content: 본문 내용
            style_config: 서식 설정 (폰트, 크기 등)
            filename: 출력 파일명
            images: 삽입할 이미지 경로 리스트 (선택)
        
        Returns:
            생성된 파일 경로
        """
        # 기본 설정 (개선된 스타일)
        base_style = {
            'font_name': '함초롬바탕',
            'font_name_english': 'HCR Batang',
            'heading_font_name': '함초롬바탕',
            'title_font_name': '함초롬바탕',
            'font_size': 11,
            'title_size': 22,
            'heading_level1_size': 16,
            'heading_level2_size': 14,
            'heading_level3_size': 13,
            'line_spacing': 1.5,
            'paragraph_spacing': 8,
            'margin_top': 2.5,
            'margin_bottom': 2.5,
            'margin_left': 2.5,
            'margin_right': 2.5,
            'treat_images_as_text': False,
            'image_placeholder_text': '※ 참고 이미지: {keyword}',
            'prepared_by': os.getenv('DOCUMENT_PREPARED_BY', 'HWP Agent AI'),
            'organization': os.getenv('DOCUMENT_ORGANIZATION', 'HWP Agent Lab')
        }
        config = base_style.copy()
        if style_config:
            config.update(style_config)
        else:
            config.update(base_style)
        primary_font = config.get('font_name', base_style['font_name'])
        config['font_name'] = primary_font
        config['font_name_english'] = config.get('font_name_english', primary_font)
        config['heading_font_name'] = primary_font
        config['title_font_name'] = primary_font
        config['treat_images_as_text'] = bool(config.get('treat_images_as_text', False))
        config['image_placeholder_text'] = config.get('image_placeholder_text', '※ 참고 이미지: {keyword}')
        
        if filename is None:
            filename = f"{title}.docx"
        
        output_path = self.output_dir / filename
        # 문서 생성 (템플릿 기반이면 템플릿을 열어서 이어 작성)
        doc = self._create_base_document(template_path)
        self.figure_counter = 1  # reset figure numbering per document
        
        # 페이지 여백 설정
        self._set_page_margins(doc, config)

        # 기본 스타일 설정
        self._set_default_style(doc, config)

        # 템플릿 사용 시 표지/메타데이터/헤더푸터 생략 (템플릿 그대로 사용)
        if not template_path:
            self._create_cover(doc, title, config)
            self._insert_metadata_block(doc, title, config)
            self._add_header_footer(doc, title, config)
        else:
            # 템플릿 모드일 경우 본문만 교체 (스타일/헤더/푸터 유지)
            self._clear_body_content(doc)

        # 본문 추가
        self._add_content(doc, content, config, images)
        
        # 저장
        doc.save(output_path)
        
        return str(output_path)

    def _create_base_document(self, template_path: Optional[str]) -> Document:
        """템플릿이 있으면 그 파일을 기반으로 문서를 생성"""
        if not template_path:
            return Document()
        try:
            candidate = Path(template_path)
            if candidate.suffix.lower() != '.docx' and candidate.exists():
                temp_docx = candidate.with_suffix('.docx')
                shutil.copyfile(candidate, temp_docx)
                doc = Document(str(temp_docx))
                temp_docx.unlink(missing_ok=True)
                return doc
            return Document(str(candidate))
        except Exception as exc:
            print(f"[DOCX] 템플릿을 열 수 없습니다. 새 문서로 대체합니다: {exc}")
            return Document()

    def _clear_body_content(self, doc: Document):
        """기존 문서의 본문 내용만 제거 (스타일, 헤더, 푸터 보존)"""
        # 단락 제거
        for paragraph in list(doc.paragraphs):
            p = paragraph._element
            p.getparent().remove(p)
            p._p = p._element = None
            
        # 표 제거
        for table in list(doc.tables):
            t = table._element
            t.getparent().remove(t)
            t._tbl = t._element = None
    
    def _set_page_margins(self, doc: Document, config: Dict[str, Any]):
        """페이지 여백 설정 (표준 문서 형식)"""
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(config.get('margin_top', 2.5))
            section.bottom_margin = Cm(config.get('margin_bottom', 2.5))
            section.left_margin = Cm(config.get('margin_left', 2.5))
            section.right_margin = Cm(config.get('margin_right', 2.5))
    
    def _set_default_style(self, doc: Document, config: Dict[str, Any]):
        """기본 스타일 설정 (한글+영문 폰트 지원)"""
        style = doc.styles['Normal']
        
        # 폰트 설정
        font = style.font
        font.name = config.get('font_name_english', 'HCR Batang')  # 영문 폰트
        font.size = Pt(config.get('font_size', 11))
        
        # 한글 폰트 설정 (중요!)
        font_name_korean = config.get('font_name', '함초롬바탕')
        if hasattr(style._element, 'rPr') and style._element.rPr is not None:
            style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name_korean)
        
        # 단락 스타일 (개선된 간격)
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = config.get('line_spacing', 1.3)
        paragraph_format.space_after = Pt(config.get('paragraph_spacing', 6))
        paragraph_format.space_before = Pt(0)
        paragraph_format.first_line_indent = Pt(0)  # 들여쓰기 없음 (현대적 스타일)
    
    def _create_cover(self, doc: Document, title: str, config: Dict[str, Any]):
        """표지 페이지 생성 (세련된 디자인)"""
        # 상단 여백 (더 넓게)
        for _ in range(6):
            doc.add_paragraph()
        
        # 제목
        title_para = doc.add_paragraph()
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(config.get('title_size', 22))
        title_run.font.name = config.get('font_name', config.get('title_font_name', config.get('heading_font_name', config.get('font_name_english', 'HCR Batang'))))
        title_run.font.color.rgb = RGBColor(30, 30, 30)  # 부드러운 검정
        
        # 한글 폰트도 적용
        if hasattr(title_run._element.rPr, 'rFonts'):
            title_run._element.rPr.rFonts.set(qn('w:eastAsia'), config.get('font_name', '함초롬바탕'))
        
        # 부제목 또는 여백
        for _ in range(2):
            doc.add_paragraph()
        
        # 페이지 나누기
        doc.add_page_break()
    
    def _add_content(self, doc: Document, content: str, config: Dict[str, Any], images: Optional[List[str]] = None):
        """본문 추가 (서식 + 이미지 텍스트 표현 지원)"""
        lines = content.split('\n')
        image_index = 0
        treat_images_as_text = bool(config.get('treat_images_as_text'))
        placeholder_only = treat_images_as_text
        tag_pattern = re.compile(r'\[gen_img\](.+?)\[/gen_img\]')

        for raw_line in lines:
            stripped = raw_line.strip()
            tag_matches = tag_pattern.findall(stripped)
            working_line = tag_pattern.sub('', stripped).strip() if tag_matches else stripped

            if not working_line and not tag_matches:
                doc.add_paragraph()
                continue

            if working_line:
                if working_line.startswith('$$') and working_line.endswith('$$') and len(working_line) > 4:
                    math_expr = working_line.strip('$')
                    self._add_math_block(doc, math_expr, config)
                    continue
                if working_line.startswith('\\[') and working_line.endswith('\\]'):
                    math_expr = working_line[2:-2].strip()
                    self._add_math_block(doc, math_expr, config)
                    continue
                if working_line.startswith('#'):
                    is_level2_heading = working_line.startswith('##') and not working_line.startswith('###')
                    self._add_heading(doc, working_line, config)
                    if not placeholder_only and is_level2_heading and images and image_index < len(images):
                        self._add_image(
                            doc,
                            images[image_index],
                            config,
                            caption=working_line.lstrip('#').strip() or None,
                        )
                        image_index += 1
                elif working_line.startswith('- ') or working_line.startswith('* '):
                    self._add_list_item(doc, working_line[2:], config)
                elif working_line[0].isdigit() and '. ' in working_line:
                    self._add_numbered_item(doc, working_line, config)
                else:
                    self._add_paragraph(doc, working_line, config)
            elif not tag_matches:
                doc.add_paragraph()

            if tag_matches:
                for keyword in tag_matches:
                    if placeholder_only:
                        self._add_image_placeholder(doc, keyword.strip(), config)
                        continue
                    if images and image_index < len(images):
                        self._add_image(
                            doc,
                            images[image_index],
                            config,
                            caption=keyword.strip(),
                        )
                        image_index += 1
                    else:
                        # 이미지가 부족하면 품격 있는 캡션으로 대체
                        self._add_image_placeholder(doc, keyword.strip(), config)

        # 본문 내 태그로 모두 소비하지 못한 이미지가 있으면 문서 말미에 추가
        if images and image_index < len(images):
            for idx in range(image_index, len(images)):
                self._add_image(doc, images[idx], config)
    
    def _add_heading(self, doc: Document, line: str, config: Dict[str, Any]):
        """제목 추가"""
        level = 0
        while line.startswith('#'):
            level += 1
            line = line[1:].strip()
        
        level = min(level, 3)  # 최대 레벨 3
        
        heading = doc.add_heading(line, level=level)
        heading_sizes = {
            1: config.get('heading_level1_size', config.get('heading_size', 16)),
            2: config.get('heading_level2_size', config.get('heading_level1_size', 15) - 1),
            3: config.get('heading_level3_size', config.get('heading_level2_size', 14) - 1)
        }
        target_font = config.get('font_name', config.get('heading_font_name', '함초롬바탕'))
        target_size = heading_sizes.get(level, heading_sizes[3])
        
        for run in heading.runs:
            run.font.name = target_font
            run.font.size = Pt(target_size)
            run.font.color.rgb = RGBColor(0, 0, 0)
            if hasattr(run._element.rPr, 'rFonts'):
                run._element.rPr.rFonts.set(qn('w:eastAsia'), target_font)
    
    def _add_list_item(self, doc: Document, text: str, config: Dict[str, Any]):
        """리스트 아이템 추가"""
        para = doc.add_paragraph(text, style='List Bullet')
        
        for run in para.runs:
            run.font.name = config.get('font_name', '함초롬바탕')
            run.font.size = Pt(config.get('font_size', 11))
    
    def _add_numbered_item(self, doc: Document, text: str, config: Dict[str, Any]):
        """번호 리스트 추가"""
        para = doc.add_paragraph(text, style='List Number')
        
        for run in para.runs:
            run.font.name = config.get('font_name', '함초롬바탕')
            run.font.size = Pt(config.get('font_size', 11))
    
    def _add_paragraph(self, doc: Document, text: str, config: Dict[str, Any]):
        """일반 문단 추가"""
        para = doc.add_paragraph()
        paragraph_format = para.paragraph_format
        paragraph_format.line_spacing = config.get('line_spacing', 1.3)
        paragraph_format.space_after = Pt(config.get('paragraph_spacing', 6) / 2)
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # 굵게/기울임 등 인라인 서식 처리
        self._add_formatted_text(para, text, config)
    
    def _add_formatted_text(self, para, text: str, config: Dict[str, Any]):
        """서식이 포함된 텍스트 추가 (한영 폰트 + 수식 표현)"""
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|__.*?__|_.*?_)', text)
        for part in parts:
            if not part:
                continue
            bold = part.startswith('**') or part.startswith('__')
            italic = (part.startswith('*') and not part.startswith('**')) or (part.startswith('_') and not part.startswith('__'))
            segment = part.strip('*_')
            self._append_math_segments(para, segment, config, bold=bold, italic=italic)

    def _append_math_segments(self, para, text: str, config: Dict[str, Any], bold: bool = False, italic: bool = False):
        if not text:
            return
        normalized = self._strip_latex_wrappers(text)
        font_name = config.get('font_name', config.get('font_name_english', 'HCR Batang'))
        pattern = re.compile(r'(\^|_)(\{([^{}]+)\}|([^\s\^_\{\}]+))')
        pos = 0
        for match in pattern.finditer(normalized):
            start, end = match.span()
            if start > pos:
                self._add_text_run(para, normalized[pos:start], font_name, config, bold, italic)
            content = match.group(3) or match.group(4) or ''
            superscript = match.group(1) == '^'
            subscript = match.group(1) == '_'
            self._add_text_run(para, content, font_name, config, bold, italic, superscript, subscript)
            pos = end
        if pos < len(normalized):
            self._add_text_run(para, normalized[pos:], font_name, config, bold, italic)

    def _add_text_run(self, para, text: str, font_name: str, config: Dict[str, Any], bold=False, italic=False, superscript=False, subscript=False):
        if not text:
            return
        run = para.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(config.get('font_size', 11))
        run.bold = bold
        run.italic = italic
        run.font.superscript = superscript
        run.font.subscript = subscript
        if hasattr(run._element.rPr, 'rFonts'):
            run._element.rPr.rFonts.set(qn('w:eastAsia'), config.get('font_name', font_name))

    def _replace_latex_symbols(self, text: str) -> str:
        if not text:
            return ''
        replaced = text
        for latex, symbol in {**LATEX_SYMBOL_MAP, **LATEX_GREEK_MAP}.items():
            replaced = replaced.replace(latex, symbol)
        return replaced

    def _add_math_block(self, doc: Document, expression: str, config: Dict[str, Any]):
        cleaned = self._strip_latex_wrappers(expression)
        para = doc.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(12)
        run = para.add_run(cleaned)
        run.font.name = config.get('font_name', '함초롬바탕')
        run.font.size = Pt(max(config.get('font_size', 11) + 2, 14))
        if hasattr(run._element.rPr, 'rFonts'):
            run._element.rPr.rFonts.set(qn('w:eastAsia'), config.get('font_name', '함초롬바탕'))

    def _strip_latex_wrappers(self, text: str) -> str:
        if not text:
            return ''
        cleaned = text.strip()
        # Remove inline/block delimiters
        for token in ('$$', '$', r'\(', r'\)', r'\[', r'\]', r'\left', r'\right'):
            cleaned = cleaned.replace(token, '')
        cleaned = re.sub(r'\\frac\{([^{}]+)\}\{([^{}]+)\}', r'(\1)/(\2)', cleaned)
        cleaned = re.sub(r'\\sqrt\{([^{}]+)\}', r'√(\1)', cleaned)
        cleaned = self._replace_latex_symbols(cleaned)
        return cleaned.strip()

    def _add_image(self, doc: Document, image_path: str, config: Dict[str, Any], caption: Optional[str] = None):
        """이미지 삽입 (자동 크기 조절 + 캡션)"""
        try:
            if not os.path.exists(image_path):
                print(f"[WARNING] Image not found: {image_path}")
                return
            
            inferred_caption = caption
            # fallback: derive from filename for professional-looking captions
            if not inferred_caption:
                stem = Path(image_path).stem.replace('_', ' ')
                inferred_caption = stem.strip() or "참고 이미지"
            
            fig_number = getattr(self, "figure_counter", None)
            label_prefix = f"그림 {fig_number}. " if fig_number is not None else ""
    
            # 이미지 단락 추가
            para = doc.add_paragraph()
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 이미지 삽입 (너비 기준 자동 조절, 최대 6인치)
            run = para.add_run()
            run.add_picture(image_path, width=Inches(4.3))
            
            # 이미지 아래 여백
            para.paragraph_format.space_after = Pt(6)

            # 캡션 (센터 정렬, 회색 톤, 이탤릭)
            caption_para = doc.add_paragraph(f"{label_prefix}{inferred_caption}")
            caption_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            caption_para.paragraph_format.space_after = Pt(12)
            for r in caption_para.runs:
                r.font.name = config.get('font_name', '함초롬바탕')
                if hasattr(r._element.rPr, 'rFonts'):
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), config.get('font_name', '함초롬바탕'))
                r.font.size = Pt(max(config.get('font_size', 11) - 1, 9))
                r.font.color.rgb = RGBColor(90, 90, 90)
                r.italic = True

            if fig_number is not None:
                self.figure_counter = fig_number + 1
            
        except Exception as e:
            print(f"[ERROR] Failed to add image {image_path}: {str(e)}")

    def _add_image_placeholder(self, doc: Document, keyword: str, config: Dict[str, Any]):
        """이미지 대신 텍스트 캡션을 추가하여 정갈한 느낌을 부여"""
        template = config.get('image_placeholder_text', '※ 참고 이미지: {keyword}') or '※ 참고 이미지: {keyword}'
        safe_keyword = keyword.strip() or '삽화'
        if '{keyword}' in template:
            placeholder = template.replace('{keyword}', safe_keyword)
        else:
            placeholder = f"{template} {safe_keyword}"

        para = doc.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after = Pt(8)
        run = para.add_run(placeholder)
        run.italic = True
        run.font.name = config.get('font_name_english', 'HCR Batang')
        run.font.size = Pt(config.get('font_size', 11))
        if hasattr(run._element.rPr, 'rFonts'):
            run._element.rPr.rFonts.set(qn('w:eastAsia'), config.get('font_name', '함초롬바탕'))

    def _insert_metadata_block(self, doc: Document, title: str, config: Dict[str, Any]):
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Light Grid Accent 1'
        labels = ['문서 제목', '작성자', '작성일']
        values = [
            title or '문서',
            config.get('prepared_by') or os.getenv('USER', 'HWP Agent AI'),
            datetime.now().strftime('%Y-%m-%d %H:%M')
        ]
        for idx, row in enumerate(table.rows):
            label_cell = row.cells[0]
            value_cell = row.cells[1]
            label_cell.text = labels[idx]
            value_cell.text = values[idx]
            for paragraph in label_cell.paragraphs:
                if paragraph.runs:
                    paragraph.runs[0].bold = True
                paragraph.style = doc.styles['Normal']
            for paragraph in value_cell.paragraphs:
                paragraph.style = doc.styles['Normal']
        doc.add_paragraph()

    def _add_header_footer(self, doc: Document, title: str, config: Dict[str, Any]):
        for section in doc.sections:
            header = section.header
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            header_para.text = f"{config.get('organization', 'HWP Agent Lab')} | {title or '문서'}"
            header_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            for run in header_para.runs:
                run.font.name = config.get('font_name', '함초롬바탕')
                run.font.size = Pt(9)
                if hasattr(run._element.rPr, 'rFonts'):
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), config.get('font_name', '함초롬바탕'))

            footer = section.footer
            para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            para.text = ''
            para.paragraph_format.space_before = Pt(6)
            run = para.add_run('Page ')
            self._add_page_field(para, 'PAGE')
            para.add_run(' / ')
            self._add_page_field(para, 'NUMPAGES')

    def _add_page_field(self, paragraph, field: str):
        run = paragraph.add_run()
        fld_char_begin = OxmlElement('w:fldChar')
        fld_char_begin.set(qn('w:fldCharType'), 'begin')
        instr_text = OxmlElement('w:instrText')
        instr_text.set(qn('xml:space'), 'preserve')
        instr_text.text = field
        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fld_char_begin)
        run._r.append(instr_text)
        run._r.append(fld_char_end)
