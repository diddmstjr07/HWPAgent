"""
HWP (워드 기반 한글 문서) 파일 생성 모듈 (함초롬바탕 버전)
- 내부는 DOCX 기반이지만 확장자만 .hwp로 저장됨
- 한글, 워드 모두에서 정상적으로 열림
- 기본 줄간 1.3배, 본문 양쪽 정렬, 이미지 중앙 정렬
- 폰트: 함초롬바탕 (HCR Batang)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from pathlib import Path
from typing import Optional, Dict, Any, List
from urllib.parse import urlparse
import html as html_lib
import importlib.util
import os
import re
import shutil
import subprocess
import sys
import zipfile
import xml.etree.ElementTree as ET
import olefile
import zlib
import struct
from bs4 import BeautifulSoup


LATEX_SYMBOL_MAP = {
    r'\times': '×',
    r'\cdot': '·',
    r'\pm': '±',
    r'\mp': '∓',
    r'\leq': '≤',
    r'\geq': '≥',
    r'\neq': '≠',
    r'\approx': '≈',
    r'\sim': '∼',
    r'\infty': '∞',
    r'\partial': '∂',
    r'\nabla': '∇',
    r'\sum': '∑',
    r'\prod': '∏',
    r'\int': '∫',
    r'\oint': '∮',
    r'\propto': '∝',
    r'\forall': '∀',
    r'\exists': '∃',
    r'\cup': '∪',
    r'\cap': '∩',
    r'\subseteq': '⊆',
    r'\supseteq': '⊇',
    r'\subset': '⊂',
    r'\supset': '⊃',
    r'\in': '∈',
    r'\notin': '∉',
    r'\perp': '⊥',
    r'\angle': '∠'
}

LATEX_GREEK_MAP = {
    r'\alpha': 'α', r'\beta': 'β', r'\gamma': 'γ', r'\delta': 'δ', r'\epsilon': 'ε',
    r'\zeta': 'ζ', r'\eta': 'η', r'\theta': 'θ', r'\iota': 'ι', r'\kappa': 'κ',
    r'\lambda': 'λ', r'\mu': 'μ', r'\nu': 'ν', r'\xi': 'ξ', r'\pi': 'π',
    r'\rho': 'ρ', r'\sigma': 'σ', r'\tau': 'τ', r'\upsilon': 'υ', r'\phi': 'φ',
    r'\chi': 'χ', r'\psi': 'ψ', r'\omega': 'ω',
    r'\Gamma': 'Γ', r'\Delta': 'Δ', r'\Theta': 'Θ', r'\Lambda': 'Λ', r'\Xi': 'Ξ',
    r'\Pi': 'Π', r'\Sigma': 'Σ', r'\Phi': 'Φ', r'\Psi': 'Ψ', r'\Omega': 'Ω'
}


class HWPHandler:
    """DOCX 기반 HWP 문서 생성기 (한글에서도 열림)"""

    def __init__(self, output_dir: str = "output"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)

    def read_hwp(self, file_path: str) -> str:
        """HWP 또는 HWPX 파일의 텍스트 내용을 추출합니다."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        # 1. Try HWPX (Zip based)
        if zipfile.is_zipfile(file_path):
            try:
                return self._read_hwpx(file_path)
            except Exception as e:
                print(f"HWPX reading failed: {e}")

        # 2. Try HWP V5 (OLE based)
        try:
            if olefile.isOleFile(file_path):
                return self._read_hwp_v5(file_path)
        except Exception as e:
            print(f"HWP V5 reading failed: {e}")

        return ""

    def _read_hwpx(self, file_path: str) -> str:
        text_content = []
        with zipfile.ZipFile(file_path, 'r') as zf:
            # Find section files
            section_files = [f for f in zf.namelist() if f.startswith('Contents/section') and f.endswith('.xml')]
            section_files.sort()

            for section_file in section_files:
                xml_data = zf.read(section_file)
                root = ET.fromstring(xml_data)

                # Extract text from <hp:t> tags
                # Use robust iterator to find 't' tags regardless of namespace prefix
                for elem in root.iter():
                    if elem.tag.endswith('}t'): # Matches <hp:t>
                        if elem.text:
                            text_content.append(elem.text)
                    elif elem.tag.endswith('}p'): # Paragraph break
                        text_content.append('\n')

        return "".join(text_content).strip()

    def _read_hwp_v5(self, file_path: str) -> str:
        # Try extracting PrvText (Preview Text)
        with olefile.OleFileIO(file_path) as ole:
            if ole.exists('PrvText'):
                # PrvText is usually UTF-16LE
                stream = ole.openstream('PrvText')
                content = stream.read().decode('utf-16le', errors='ignore')
                return content

            return "[HWP V5: 텍스트 미리보기(PrvText) 스트림이 없습니다. 텍스트 추출이 제한됩니다.]"

    def convert_hwp_to_html(self, file_path: str, template_id: Optional[str] = None) -> str:
        """
        HWP 파일을 HTML로 변환합니다.
        
        Args:
            file_path (str): HWP 파일 경로

        Returns:
            str: 변환된 HTML 문자열
        """
        def wrap_text_as_html(text: str) -> str:
            escaped = html_lib.escape(text or "")
            paragraphs = []
            buffer = []
            for line in escaped.splitlines():
                if line.strip():
                    buffer.append(line)
                else:
                    if buffer:
                        paragraphs.append("<p>" + "<br>".join(buffer) + "</p>")
                        buffer = []
            if buffer:
                paragraphs.append("<p>" + "<br>".join(buffer) + "</p>")
            return "\n".join(paragraphs) if paragraphs else "<p></p>"

        command = self._resolve_hwp5html_command()
        if not command:
            print("[HWP HTML] hwp5html not available. Falling back to text extraction.")
            text_content = self.read_hwp(file_path)
            return wrap_text_as_html(text_content)

        safe_id = template_id or Path(file_path).stem
        html_root = self.output_dir / "templates_html"
        html_root.mkdir(parents=True, exist_ok=True)
        output_dir = html_root / safe_id

        if output_dir.exists():
            shutil.rmtree(output_dir)

        try:
            subprocess.run(
                [*command, "--output", str(output_dir), file_path],
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True
            )
        except subprocess.CalledProcessError as exc:
            err_msg = exc.stderr.strip() if exc.stderr else str(exc)
            print(f"[HWP HTML] hwp5html failed: {err_msg}")
            text_content = self.read_hwp(file_path)
            return wrap_text_as_html(text_content)
        except Exception as exc:
            print(f"[HWP HTML] hwp5html error: {exc}")
            text_content = self.read_hwp(file_path)
            return wrap_text_as_html(text_content)

        index_path = output_dir / "index.xhtml"
        if not index_path.exists():
            index_path = output_dir / "index.html"

        if not index_path.exists():
            text_content = self.read_hwp(file_path)
            return wrap_text_as_html(text_content)

        html_text = index_path.read_text(encoding="utf-8", errors="ignore")
        soup = BeautifulSoup(html_text, "html.parser")
        if not soup.body:
            text_content = self.read_hwp(file_path)
            return wrap_text_as_html(text_content)

        asset_base = f"/api/template/asset/{safe_id}/"
        for img in soup.find_all("img"):
            src = img.get("src")
            if not src:
                continue
            parsed = urlparse(src)
            if parsed.scheme or src.startswith(("data:", "blob:")):
                continue
            img["src"] = asset_base + src.lstrip("/")

        for link in soup.find_all("link"):
            href = link.get("href")
            if not href:
                continue
            parsed = urlparse(href)
            if parsed.scheme or href.startswith(("data:", "blob:")):
                continue
            link["href"] = asset_base + href.lstrip("/")

        return str(soup)

    def _resolve_hwp5html_command(self) -> Optional[List[str]]:
        env_override = os.getenv("HWP5HTML_PATH")
        if env_override:
            candidate = Path(env_override).expanduser()
            if candidate.exists():
                return [str(candidate)]

        hwp5html_path = shutil.which("hwp5html")
        if hwp5html_path:
            return [hwp5html_path]

        exe_name = "hwp5html.exe" if os.name == "nt" else "hwp5html"
        venv_prefix = os.getenv("VIRTUAL_ENV")
        conda_prefix = os.getenv("CONDA_PREFIX")
        for prefix in (venv_prefix, conda_prefix):
            if not prefix:
                continue
            bin_dir = "Scripts" if os.name == "nt" else "bin"
            candidate = Path(prefix) / bin_dir / exe_name
            if candidate.exists():
                return [str(candidate)]

        candidate = Path(sys.executable).resolve().parent / exe_name
        if candidate.exists():
            return [str(candidate)]

        try:
            if importlib.util.find_spec("hwp5.hwp5html"):
                return [sys.executable, "-m", "hwp5.hwp5html"]
        except (ModuleNotFoundError, ValueError):
            pass

        try:
            if importlib.util.find_spec("hwp5html"):
                return [sys.executable, "-m", "hwp5html"]
        except (ModuleNotFoundError, ValueError):
            pass

        return None

    def create_hwp_document(
        self,
        title: str,
        content: str,
        style_config: Optional[Dict[str, Any]] = None,
        filename: Optional[str] = None,
        images: Optional[List[str]] = None
    ) -> str:
        """HWP 문서 생성 (DOCX 구조 기반, .hwp 확장자)"""
        base_style = {
            'font_name': '함초롬바탕',
            'font_name_english': 'Times New Roman',
            'heading_font_name': '함초롬바탕',
            'title_font_name': '함초롬바탕',
            'font_size': 11,
            'title_size': 22,
            'heading_level1_size': 16,
            'heading_level2_size': 14,
            'heading_level3_size': 13,
            'line_spacing': 1.3,
            'paragraph_spacing': 6,
            'margin_top': 2.5,
            'margin_bottom': 2.5,
            'margin_left': 2.5,
            'margin_right': 2.5
        }

        style = base_style.copy()
        if style_config:
            style.update(style_config)
        style['heading_font_name'] = style.get('heading_font_name', style.get('font_name', base_style['font_name']))
        style['title_font_name'] = style.get('title_font_name', style['heading_font_name'])

        if filename is None:
            filename = f"{title}.hwp"
        elif not filename.endswith(".hwp"):
            filename += ".hwp"

        output_path = self.output_dir / filename

        doc = Document()
        self.figure_counter = 1  # reset numbering per document
        self._set_page_margins(doc, style)
        self._set_default_style(doc, style)

        self._create_cover(doc, title, style)
        self._add_content(doc, content, style, images)

        temp_path = output_path.with_suffix(".docx")
        doc.save(temp_path)
        temp_path.rename(output_path)

        print(f"✅ HWP 문서 생성 완료: {output_path}")
        return str(output_path)

    # --------------------------------------------------------------
    # 내부 메서드
    # --------------------------------------------------------------
    def _set_page_margins(self, doc: Document, config: Dict[str, Any]):
        for section in doc.sections:
            section.top_margin = Cm(config.get('margin_top', 2.5))
            section.bottom_margin = Cm(config.get('margin_bottom', 2.5))
            section.left_margin = Cm(config.get('margin_left', 2.5))
            section.right_margin = Cm(config.get('margin_right', 2.5))

    def _set_default_style(self, doc: Document, config: Dict[str, Any]):
        style = doc.styles['Normal']
        font = style.font
        font.name = config.get('font_name_english', 'Times New Roman')
        font.size = Pt(config.get('font_size', 11))

        font_name_korean = config.get('font_name', '함초롬바탕')
        if hasattr(style._element, 'rPr') and style._element.rPr is not None:
            style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name_korean)

        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = config.get('line_spacing', 1.3)
        paragraph_format.space_after = Pt(config.get('paragraph_spacing', 6))
        paragraph_format.space_before = Pt(0)
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def _create_cover(self, doc: Document, title: str, config: Dict[str, Any]):
        for _ in range(6):
            doc.add_paragraph()

        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(config.get('title_size', 22))
        title_font = config.get('font_name', config.get('title_font_name', config.get('heading_font_name', config.get('font_name_english', 'Times New Roman'))))
        title_run.font.name = title_font
        title_run.font.color.rgb = RGBColor(20, 20, 20)
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), title_font)

        subtitle = doc.add_paragraph("기술 보고서 / Research Report")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run()
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(100, 100, 100)
        run.font.name = title_font
        run._element.rPr.rFonts.set(qn('w:eastAsia'), title_font)

        for _ in range(2):
            doc.add_paragraph()
        doc.add_page_break()

    def _add_content(self, doc: Document, content: str, config: Dict[str, Any], images: Optional[List[str]] = None):
        lines = content.split('\n')
        image_index = 0
        for i, line in enumerate(lines):
            original_line = line
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue
            if line.startswith('#'):
                is_level2_heading = original_line.strip().startswith('##') and not original_line.strip().startswith('###')
                self._add_heading(doc, line, config)
                if images and image_index < len(images) and is_level2_heading:
                    caption = line.lstrip('#').strip() or None
                    self._add_image(doc, images[image_index], config, caption=caption)
                    image_index += 1
            elif line.startswith('- ') or line.startswith('* '):
                self._add_list_item(doc, line[2:], config)
            elif line[0].isdigit() and '. ' in line:
                self._add_numbered_item(doc, line, config)
            elif (line.startswith('$$') and line.endswith('$$') and len(line) > 4) or (line.startswith('\\[') and line.endswith('\\]')):
                math_expr = line.strip('$')
                if math_expr.startswith('\\[') and math_expr.endswith('\\]'):
                    math_expr = math_expr[2:-2].strip()
                self._add_math_block(doc, math_expr, config)
            else:
                self._add_paragraph(doc, line, config)

    def _add_heading(self, doc: Document, line: str, config: Dict[str, Any]):
        level = line.count('#')
        text = line.strip('#').strip()
        heading = doc.add_heading(text, level=min(level, 3))
        heading.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        heading_sizes = {
            1: config.get('heading_level1_size', config.get('heading_size', 16)),
            2: config.get('heading_level2_size', config.get('heading_level1_size', 15) - 1),
            3: config.get('heading_level3_size', config.get('heading_level2_size', 14) - 1)
        }
        target_font = config.get('heading_font_name', config.get('font_name', '함초롬바탕'))
        target_size = heading_sizes.get(level, heading_sizes[3])
        for run in heading.runs:
            run.font.name = target_font
            run.font.size = Pt(target_size)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), target_font)
            run.font.color.rgb = RGBColor(0, 0, 0)

    def _add_list_item(self, doc: Document, text: str, config: Dict[str, Any]):
        para = doc.add_paragraph(text, style='List Bullet')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in para.runs:
            run.font.name = config.get('font_name', '함초롬바탕')
            run._element.rPr.rFonts.set(qn('w:eastAsia'), config.get('font_name', '함초롬바탕'))
            run.font.size = Pt(config.get('font_size', 11))

    def _add_numbered_item(self, doc: Document, text: str, config: Dict[str, Any]):
        para = doc.add_paragraph(text, style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in para.runs:
            run.font.name = config.get('font_name', '함초롬바탕')
            run._element.rPr.rFonts.set(qn('w:eastAsia'), config.get('font_name', '함초롬바탕'))
            run.font.size = Pt(config.get('font_size', 11))

    def _add_paragraph(self, doc: Document, text: str, config: Dict[str, Any]):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.line_spacing = config.get('line_spacing', 1.3)
        para.paragraph_format.space_after = Pt(config.get('paragraph_spacing', 6))
        self._add_formatted_text(para, text, config)

    def _add_formatted_text(self, para, text: str, config: Dict[str, Any]):
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|__.*?__|_.*?_)', text)
        for part in parts:
            if not part:
                continue
            bold = part.startswith('**') or part.startswith('__')
            italic = (part.startswith('*') and not part.startswith('**')) or (part.startswith('_') and not part.startswith('__'))
            segment = part.strip('*_')
            self._append_math_segments(para, segment, config, bold=bold, italic=italic)

    def _append_math_segments(self, para, text: str, config: Dict[str, Any], bold: bool = False, italic: bool = False):
        if not text:
            return
        normalized = self._strip_latex_wrappers(text)
        font_name = config.get('font_name', config.get('font_name_english', 'Times New Roman'))
        pattern = re.compile(r'(\^|_)(\{([^{}]+)\}|([^\s\^_\{\}]+))')
        pos = 0
        for match in pattern.finditer(normalized):
            start, end = match.span()
            if start > pos:
                self._add_text_run(para, normalized[pos:start], font_name, config, bold, italic)
            content = match.group(3) or match.group(4) or ''
            superscript = match.group(1) == '^'
            subscript = match.group(1) == '_'
            self._add_text_run(para, content, font_name, config, bold, italic, superscript, subscript)
            pos = end
        if pos < len(normalized):
            self._add_text_run(para, normalized[pos:], font_name, config, bold, italic)

    def _add_text_run(
        self,
        para,
        text: str,
        font_name: str,
        config: Dict[str, Any],
        bold: bool = False,
        italic: bool = False,
        superscript: bool = False,
        subscript: bool = False,
    ):
        if not text:
            return
        run = para.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(config.get('font_size', 11))
        run.bold = bold
        run.italic = italic
        run.font.superscript = superscript
        run.font.subscript = subscript
        if hasattr(run._element.rPr, 'rFonts'):
            run._element.rPr.rFonts.set(qn('w:eastAsia'), config.get('font_name', font_name))

    def _replace_latex_symbols(self, text: str) -> str:
        replaced = text
        for latex, symbol in {**LATEX_SYMBOL_MAP, **LATEX_GREEK_MAP}.items():
            replaced = replaced.replace(latex, symbol)
        return replaced

    def _strip_latex_wrappers(self, text: str) -> str:
        if not text:
            return ''
        cleaned = text.strip()
        for token in ('$$', '$', r'\(', r'\)', r'\[', r'\]', r'\left', r'\right'):
            cleaned = cleaned.replace(token, '')
        cleaned = re.sub(r'\\frac\{([^{}]+)\}\{([^{}]+)\}', r'(\1)/(\2)', cleaned)
        cleaned = re.sub(r'\\sqrt\{([^{}]+)\}', r'√(\1)', cleaned)
        cleaned = self._replace_latex_symbols(cleaned)
        return cleaned.strip()

    def _add_math_block(self, doc: Document, expression: str, config: Dict[str, Any]):
        cleaned = self._strip_latex_wrappers(expression)
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(12)
        run = para.add_run(cleaned)
        run.font.name = config.get('font_name', '함초롬바탕')
        run.font.size = Pt(max(config.get('font_size', 11) + 2, 14))
        if hasattr(run._element.rPr, 'rFonts'):
            run._element.rPr.rFonts.set(qn('w:eastAsia'), config.get('font_name', '함초롬바탕'))

    def _add_image(self, doc: Document, image_path: str, config: Dict[str, Any], caption: Optional[str] = None):
        try:
            if not os.path.exists(image_path):
                print(f"[WARNING] Image not found: {image_path}")
                return
            para = doc.add_paragraph()
            run = para.add_run()
            run.add_picture(image_path, width=Inches(4.3))
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.line_spacing = config.get('line_spacing', 1.3)

            inferred_caption = caption
            if inferred_caption is None:
                stem = Path(image_path).stem.replace('_', ' ')
                inferred_caption = stem.strip() or "참고 이미지"
            fig_number = getattr(self, "figure_counter", None)
            label_prefix = f"그림 {fig_number}. " if fig_number is not None else ""

            caption_para = doc.add_paragraph(f"{label_prefix}{inferred_caption}")
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_para.paragraph_format.space_after = Pt(10)
            for r in caption_para.runs:
                r.font.name = config.get('font_name', '함초롬바탕')
                r._element.rPr.rFonts.set(qn('w:eastAsia'), config.get('font_name', '함초롬바탕'))
                r.font.size = Pt(max(config.get('font_size', 11) - 1, 9))
                r.font.color.rgb = RGBColor(90, 90, 90)
                r.italic = True

            if fig_number is not None:
                self.figure_counter = fig_number + 1
        except Exception as e:
            print(f"[ERROR] Failed to add image {image_path}: {str(e)}")


# --------------------------------------------------------------
# 실행 예시
# --------------------------------------------------------------
if __name__ == "__main__":
    handler = HWPHandler()
    handler.create_hwp_document(
        title="Aphonia 프로젝트 보고서",
        content="# 개요\n\n본 문서는 Aphonia 시스템에 대한 기술 보고서입니다.\n\n## 실험 결과\n\n- EMG 데이터 처리\n- Lip-reading 기반 예측\n\n## 결론\n\n본 연구는 무성증 환자의 의사소통 가능성을 확장합니다.",
        images=["output/images/coral_reef.jpg"],
        filename="aphonia_report.hwp",
    )
